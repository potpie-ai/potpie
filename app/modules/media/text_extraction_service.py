"""Text extraction service for document attachments."""
import logging
from typing import Optional, Tuple
from io import BytesIO
import chardet

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Base exception for text extraction errors."""
    pass


class TextExtractionService:
    """Handles text extraction from various document formats."""

    # Size threshold for storing in JSONB vs S3 (50KB)
    INLINE_STORAGE_THRESHOLD = 50 * 1024

    def __init__(self):
        # Lazy import heavy libraries
        self._pypdf = None
        self._docx = None
        self._pandas = None

    @property
    def pypdf(self):
        """Lazy load PyPDF2."""
        if self._pypdf is None:
            try:
                import pypdf
                self._pypdf = pypdf
            except ImportError as e:
                raise TextExtractionError("pypdf library not installed") from e
        return self._pypdf

    @property
    def docx(self):
        """Lazy load python-docx."""
        if self._docx is None:
            try:
                import docx
                self._docx = docx
            except ImportError as e:
                raise TextExtractionError("python-docx library not installed") from e
        return self._docx

    @property
    def pandas(self):
        """Lazy load pandas."""
        if self._pandas is None:
            try:
                import pandas as pd
                self._pandas = pd
            except ImportError as e:
                raise TextExtractionError("pandas library not installed") from e
        return self._pandas

    def extract_text(
        self,
        file_data: bytes,
        mime_type: str,
        file_name: str
    ) -> Tuple[str, dict]:
        """
        Extract text from file data.

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            if "pdf" in mime_type:
                return self._extract_pdf(file_data)
            elif "word" in mime_type or "vnd.openxmlformats-officedocument.wordprocessingml" in mime_type:
                return self._extract_docx(file_data)
            elif "csv" in mime_type or mime_type == "text/csv":
                return self._extract_csv(file_data)
            elif "spreadsheet" in mime_type or "vnd.openxmlformats-officedocument.spreadsheetml" in mime_type:
                return self._extract_xlsx(file_data)
            elif mime_type == "text/plain" or mime_type.startswith("text/"):
                return self._extract_text_file(file_data, mime_type)
            elif self._is_code_file(file_name, mime_type):
                return self._extract_code_file(file_data)
            else:
                raise TextExtractionError(f"Unsupported file type: {mime_type}")

        except TextExtractionError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error extracting text: {e}", exc_info=True)
            raise TextExtractionError(f"Failed to extract text: {str(e)}") from e

    def _extract_pdf(self, file_data: bytes) -> Tuple[str, dict]:
        """Extract text from PDF."""
        try:
            pdf_file = BytesIO(file_data)
            reader = self.pypdf.PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")

            extracted_text = "\n\n".join(text_parts)

            metadata = {
                "extraction_method": "pypdf",
                "page_count": len(reader.pages),
                "pdf_metadata": dict(reader.metadata) if reader.metadata else {},
            }

            return extracted_text, metadata

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise TextExtractionError(f"Failed to extract PDF: {str(e)}") from e

    def _extract_docx(self, file_data: bytes) -> Tuple[str, dict]:
        """Extract text from DOCX."""
        try:
            docx_file = BytesIO(file_data)
            doc = self.docx.Document(docx_file)

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Extract tables
            table_texts = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    table_texts.append("\n".join(table_data))

            extracted_text = "\n\n".join(paragraphs)
            if table_texts:
                extracted_text += "\n\n=== TABLES ===\n\n" + "\n\n".join(table_texts)

            metadata = {
                "extraction_method": "python-docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            }

            return extracted_text, metadata

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise TextExtractionError(f"Failed to extract DOCX: {str(e)}") from e

    def _extract_csv(self, file_data: bytes) -> Tuple[str, dict]:
        """Extract text from CSV."""
        try:
            # Try UTF-8 first, then detect encoding
            try:
                csv_text = file_data.decode('utf-8')
            except UnicodeDecodeError:
                detected = chardet.detect(file_data)
                encoding = detected['encoding'] or 'latin-1'
                csv_text = file_data.decode(encoding)

            # Parse with pandas for better formatting
            import io
            df = self.pandas.read_csv(io.StringIO(csv_text))

            # Convert to readable format
            extracted_text = f"CSV Data ({len(df)} rows, {len(df.columns)} columns):\n\n"
            extracted_text += df.to_string(index=False)

            metadata = {
                "extraction_method": "pandas",
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": list(df.columns),
            }

            return extracted_text, metadata

        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            raise TextExtractionError(f"Failed to extract CSV: {str(e)}") from e

    def _extract_xlsx(self, file_data: bytes) -> Tuple[str, dict]:
        """Extract text from XLSX."""
        try:
            xlsx_file = BytesIO(file_data)

            # Read all sheets
            excel_file = self.pandas.ExcelFile(xlsx_file)
            sheet_texts = []

            for sheet_name in excel_file.sheet_names:
                df = self.pandas.read_excel(excel_file, sheet_name=sheet_name)
                sheet_text = f"=== Sheet: {sheet_name} ===\n"
                sheet_text += f"({len(df)} rows, {len(df.columns)} columns)\n\n"
                sheet_text += df.to_string(index=False)
                sheet_texts.append(sheet_text)

            extracted_text = "\n\n".join(sheet_texts)

            metadata = {
                "extraction_method": "pandas",
                "sheet_count": len(excel_file.sheet_names),
                "sheet_names": excel_file.sheet_names,
            }

            return extracted_text, metadata

        except Exception as e:
            logger.error(f"XLSX extraction failed: {e}")
            raise TextExtractionError(f"Failed to extract XLSX: {str(e)}") from e

    def _extract_text_file(self, file_data: bytes, mime_type: str) -> Tuple[str, dict]:
        """Extract text from plain text file."""
        try:
            # Detect encoding
            detected = chardet.detect(file_data)
            encoding = detected['encoding'] or 'utf-8'

            extracted_text = file_data.decode(encoding)

            metadata = {
                "extraction_method": "decode",
                "encoding": encoding,
                "confidence": detected['confidence'],
            }

            return extracted_text, metadata

        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            raise TextExtractionError(f"Failed to extract text file: {str(e)}") from e

    def _is_code_file(self, file_name: str, mime_type: str) -> bool:
        """Check if file is a code file based on extension."""
        code_extensions = {
            '.py', '.js', '.ts', '.tsx', '.jsx', '.java', '.cpp', '.c', '.h',
            '.cs', '.rb', '.go', '.rs', '.php', '.swift', '.kt', '.scala',
            '.sh', '.bash', '.sql', '.r', '.m', '.mm', '.md', '.json', '.xml',
            '.yaml', '.yml', '.toml', '.ini', '.conf', '.cfg'
        }

        extension = '.' + file_name.split('.')[-1].lower() if '.' in file_name else ''
        return extension in code_extensions or 'application/json' in mime_type

    def _extract_code_file(self, file_data: bytes) -> Tuple[str, dict]:
        """Extract text from code file."""
        # Same as text file but note it's code
        extracted_text, metadata = self._extract_text_file(file_data, "text/plain")
        metadata["extraction_method"] = "code_decode"
        metadata["is_code"] = True
        return extracted_text, metadata

    def should_store_inline(self, extracted_text: str) -> bool:
        """Determine if extracted text should be stored inline in JSONB."""
        return len(extracted_text.encode('utf-8')) < self.INLINE_STORAGE_THRESHOLD
